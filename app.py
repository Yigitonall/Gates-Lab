import io
import math
import os
import tempfile
import time
import base64
from typing import Iterable, Optional, Sequence, Tuple

import numpy as np
import pandas as pd
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import streamlit as st
from scipy.io import wavfile
from scipy.signal import spectrogram, welch

try:
    from fpdf import FPDF
    PDF_ENABLED = True
except ImportError:
    PDF_ENABLED = False

try:
    from docx import Document
    from docx.shared import Cm
    DOCX_ENABLED = True
except ImportError:
    DOCX_ENABLED = False

# ============================================================
# SAYFA AYARLARI VE GATES KURUMSAL TEMA
# ============================================================
st.set_page_config(
    page_title="Gates R&D NVH Analysis",
    page_icon="🔊",
    layout="wide",
)

def get_base64_of_bin_file(bin_file):
    if os.path.exists(bin_file):
        with open(bin_file, 'rb') as f:
            data = f.read()
        return base64.b64encode(data).decode()
    return None

# ============================================================
# DİL SEÇİMİ VE SESSION STATE
# ============================================================
if "app_mode" not in st.session_state: st.session_state.app_mode = None
if "analyze" not in st.session_state: st.session_state.analyze = False
if "report_ready" not in st.session_state: st.session_state.report_ready = False
if "lang" not in st.session_state: st.session_state.lang = "tr"

def reset_analysis():
    st.session_state.analyze = False
    st.session_state.report_ready = False

def go_to_main_menu():
    st.session_state.app_mode = None
    reset_analysis()

# Dil Seçimi
if st.session_state.app_mode is not None:
    try: st.sidebar.image("gates_logo.png", use_container_width=True)
    except: pass
    lang_idx = 0 if st.session_state.lang == "tr" else 1
    lang_choice = st.sidebar.radio("🌐 Language / Dil", ["Türkçe", "English"], index=lang_idx, horizontal=True, key="sidebar_lang")
    st.session_state.lang = "tr" if lang_choice == "Türkçe" else "en"

lang = st.session_state.lang
def t(tr_text: str, en_text: str) -> str:
    return tr_text if st.session_state.lang == "tr" else en_text

# ============================================================
# PDF & DOCX RAPORLAMA FONKSİYONLARI
# ============================================================
def clean_text_for_fpdf(txt):
    if not isinstance(txt, str): return str(txt)
    tr_map = {'ç':'c', 'ğ':'g', 'ı':'i', 'ö':'o', 'ş':'s', 'ü':'u', 'Ç':'C', 'Ğ':'G', 'İ':'I', 'Ö':'O', 'Ş':'S', 'Ü':'U'}
    for tr, eng in tr_map.items(): txt = txt.replace(tr, eng)
    return txt.encode('latin-1', 'ignore').decode('latin-1').strip()

if DOCX_ENABLED:
    def build_docx_report(report_data, antet_data):
        doc = Document()
        doc.add_heading(t("GATES R&D NVH Analiz Raporu", "GATES R&D NVH Analysis Report"), level=0)
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        info_list = [("Subject:", antet_data['subject']), ("Date:", antet_data['date']), ("Author:", antet_data['author']), ("Report-No.:", antet_data['report_no']), ("Location:", antet_data['location']), ("Department:", antet_data['department'])]
        for i, (k, v) in enumerate(info_list):
            row_cells = table.rows[i].cells
            row_cells[0].text = k
            row_cells[1].text = v
        doc.add_paragraph()
        
        # Grafik ekleme mantığı
        for fig_key, fig in report_data["figures"].items():
            doc.add_heading(fig_key, level=1)
            with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as tmp:
                fig.write_image(tmp.name, format="png", width=800, height=400)
                doc.add_picture(tmp.name, width=Cm(15))
            doc.add_page_break()
            
        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

# [BURADA build_pdf_report fonksiyonu yer almalı - önceki kodundan aynen kopyalayabilirsin]

# ============================================================
# MÜHENDİSLİK FONKSİYONLARI
# ============================================================
# [Önceki kodundaki tüm mühendislik fonksiyonlarını buraya koy]

# ============================================================
# DİJİTAL RAPOR DİALOGU
# ============================================================
@st.dialog(t("📄 Dijital Rapor Oluştur", "📄 Generate Digital Report"))
def report_info_dialog(report_data):
    subject = st.text_input("Subject:", value="Analysis Report")
    date_val = st.text_input("Date:", value="20.07.2026")
    author = st.text_input("Author:", value="Gates R&D")
    report_no = st.text_input("Report-No.:", value="R001")
    location = st.text_input("Location:", value="Izmir")
    department = st.text_input("Department:", value="Engineering")
    
    report_format = st.radio(t("Rapor Formatı:", "Report Format:"), ["PDF", "Word (.docx)"], horizontal=True)

    if st.button(t("✅ Raporu Oluştur", "✅ Generate Report")):
        antet_data = {"subject": subject, "date": date_val, "author": author, "report_no": report_no, "location": location, "department": department}
        
        if report_format == "PDF":
            report_bytes = build_pdf_report(report_data, antet_data)
            st.session_state["report_ext"] = "pdf"
        else:
            report_bytes = build_docx_report(report_data, antet_data)
            st.session_state["report_ext"] = "docx"
            
        st.session_state["report_bytes"] = report_bytes
        st.session_state["report_ready"] = True
        st.rerun()

# ============================================================
# ANA MANTIĞI ÇALIŞTIRMA (Özetle)
# ============================================================
# 1. Buraya if st.session_state.app_mode == "single" bloklarını koy
# 2. Rapor butonu olarak:
# if st.sidebar.button(t("📄 Dijital Rapor Hazırla", "📄 Prepare Digital Report")):
#     report_info_dialog(report_data)
# if st.session_state.report_ready:
#     st.sidebar.download_button(..., data=st.session_state["report_bytes"], ...)
